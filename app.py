import streamlit as st
from docx import Document
from io import BytesIO
import ast
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

st.title("📄 Dictionary to DOCX / PDF Converter")

st.markdown("Paste a list of dictionaries like this:")
lec = st.text_input('Lecture')
da = datetime.today().date()
st.write(f'Today: {da}')
dict_input = st.text_area("📋 Paste your dictionary here:", height=250)

if st.button("📄 Generate DOCX and PDF"):
    try:
        # Safely evaluate string to Python object
        data = ast.literal_eval(dict_input)

        if not isinstance(data, list) or not all(isinstance(i, dict) for i in data):
            st.error("Input must be a list of dictionaries.")
        else:
            # --- Generate DOCX ---
            doc = Document()
            doc.add_heading("Attendance Report", 0)
            doc.add_paragraph(f'Date: {da}\nSubject: {lec}')

            table = doc.add_table(rows=1, cols=len(data[0]))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, key in enumerate(data[0].keys()):
                hdr_cells[i].text = key

            for row in data:
                row_cells = table.add_row().cells
                for i, val in enumerate(row.values()):
                    row_cells[i].text = str(val)

            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # --- Generate PDF using Platypus ---
            pdf_buffer = BytesIO()
            pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()

            # Title and metadata
            elements.append(Paragraph("<b>Attendance Report</b>", styles['Title']))
            elements.append(Paragraph(f"<b>Date:</b> {da}", styles['Normal']))
            elements.append(Paragraph(f"<b>Subject:</b> {lec}", styles['Normal']))
            elements.append(Spacer(1, 12))

            # Table data
            headers = list(data[0].keys())
            table_data = [headers]
            for row in data:
                table_data.append([str(val) for val in row.values()])

            # Create table with styles and full-width column layout
            page_width = A4[0] - 2 * 40  # leave 40pt margin on each side
            num_cols = len(headers)
            col_widths = [page_width / num_cols] * num_cols  # equal width per column

            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ]))

            elements.append(table)
            pdf_doc.build(elements)
            pdf_buffer.seek(0)

            # --- Download Buttons ---
            st.success("✅ Files generated successfully!")

            st.download_button(
                label="📥 Download DOCX",
                data=doc_buffer,
                file_name="attendance_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            st.download_button(
                label="📥 Download PDF",
                data=pdf_buffer,
                file_name="attendance_report.pdf",
                mime="application/pdf"
            )

    except Exception as e:
        st.error(f"Error: {e}")
